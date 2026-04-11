"""Build the Pratibmb project journal as a DOCX file."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUT = os.path.join(os.path.dirname(__file__), "project_journal.docx")

ENTRIES = [
    {
        "day": "Day 1",
        "title": "Model Selection \u2014 Blind Evaluation",
        "body": (
            "Tested 6 local LLMs in a blind evaluation for persona simulation quality. "
            "The user scored responses without knowing which model produced them.\n\n"
            "Models tested:\n"
            "  \u2022 Gemma-3-4B-Instruct (2.3 GB)\n"
            "  \u2022 Gemma-2-9B-Instruct (5.5 GB)\n"
            "  \u2022 Llama-3.2-3B-Instruct (1.8 GB)\n"
            "  \u2022 Llama-3.1-8B-Instruct (4.7 GB)\n"
            "  \u2022 Mistral-7B-Instruct (4.1 GB)\n"
            "  \u2022 Qwen2.5-7B-Instruct (4.4 GB)\n\n"
            "Result: Gemma-3-4B won despite being the smallest model. "
            "It produced the most natural texting style, best voice matching, and "
            "stayed in character most consistently.\n\n"
            "Key insight: bigger is not always better for persona simulation. "
            "A smaller model that follows instructions precisely outperforms "
            "larger models that are more \"creative\" but less controllable."
        ),
    },
    {
        "day": "Day 1",
        "title": "Automated Head-to-Head: Gemma-3-4B vs Gemma-4-E4B",
        "body": (
            "Ran an automated head-to-head evaluation using Qwen2.5-7B as an "
            "LLM-as-judge. Compared Gemma-3-4B (2.3 GB) against the newer "
            "Gemma-4-E4B (5.4 GB) with A/B position randomization to eliminate "
            "position bias.\n\n"
            "Result: Gemma-3-4B won again. The newer, larger model was more "
            "verbose and tended to break character, while Gemma-3-4B maintained "
            "a consistent, casual texting persona.\n\n"
            "Decision: Ship with Gemma-3-4B-Instruct Q4_K_M as the chat model."
        ),
    },
    {
        "day": "Day 2",
        "title": "App Architecture & Scaffolding",
        "body": (
            "Designed and scaffolded the full application:\n\n"
            "Stack:\n"
            "  \u2022 Python core (importers, RAG, LLM, voice fingerprinting)\n"
            "  \u2022 Tauri 2.x desktop shell (Rust backend + HTML/CSS/JS frontend)\n"
            "  \u2022 llama-cpp-python with Metal acceleration\n"
            "  \u2022 SQLite for corpus + embeddings (float32 BLOBs)\n"
            "  \u2022 nomic-embed-text-v1.5 Q4_K_M for embeddings (84 MB)\n\n"
            "Importers: WhatsApp (.txt), Facebook DYI (JSON), Instagram DYI (JSON)\n\n"
            "Created GitHub repo (tapaskar/Pratibmb, private). AGPLv3 license with "
            "open-core monetization model. 20 unit tests all passing.\n\n"
            "The name \"Pratibmb\" is a deliberately coined spelling (not Pratibimb). "
            "Trademark research initiated."
        ),
    },
    {
        "day": "Day 2",
        "title": "Facebook Data Import \u2014 11,676 Messages",
        "body": (
            "Challenge: WhatsApp bulk export wasn't feasible. The Mac is remote "
            "(no cable for iPhone backup), not enough disk for local backup, and "
            "iCloud backup is E2E encrypted.\n\n"
            "Pivot: Used Facebook \"Download Your Information\" (DYI) JSON export.\n\n"
            "Bug found: Facebook changed their DYI folder from "
            "\"your_activity_across_facebook\" to \"your_facebook_activity\" \u2014 "
            "importer returned 0 messages until the new path was added.\n\n"
            "Result: 11,676 messages imported, 4,009 self-authored messages, "
            "spanning 2010\u20132026. Facebook's mojibake (double-encoded UTF-8 as "
            "latin-1) handled automatically."
        ),
    },
    {
        "day": "Day 2",
        "title": "Voice Fingerprinting",
        "body": (
            "Analyzed the real corpus to extract communication style:\n\n"
            "  \u2022 96% lowercase\n"
            "  \u2022 Average 5.2 words per message\n"
            "  \u2022 76% short messages (under 5 words)\n"
            "  \u2022 Hinglish detected: hai, mein, nahi, abt, yaar\n"
            "  \u2022 Occasional emojis: \U0001f60a \U0001f600 \U0001f37e\n\n"
            "Generated a voice directive injected into the system prompt to guide "
            "the LLM toward matching this style."
        ),
    },
    {
        "day": "Day 2\u20133",
        "title": "The Prompt Engineering Battle \u2014 8 Iterations",
        "body": (
            "The hardest part: getting Gemma-3-4B to actually sound like the user.\n\n"
            "Three problems:\n"
            "  1. Hallucination: Model invented \"freelance graphic designer in Brooklyn\"\n"
            "  2. Verbosity: 100+ word responses when user averages 5.2 words\n"
            "  3. Wrong language: formal English instead of Hinglish\n\n"
            "Iterations:\n"
            "  v1\u2013v3: Long instruction-heavy prompts. Model ignored all of them.\n"
            "  v4: \"ABSOLUTE RULES\" with examples. Still hallucinated.\n"
            "  v5: Minimal prompt. Too short, no personality.\n"
            "  v6: \"Continue this text conversation\" framing. Breakthrough!\n"
            "  v7: Multi-turn style priming. Backfired \u2014 model started saying \"I'm an AI.\"\n"
            "  v8: Final version. Concise framing + random seed + anti-hallucination guard.\n\n"
            "Before (v1): \"That's a really good question. Reflecting on these messages, "
            "it's complicated! I'm working as a freelance graphic designer in Brooklyn...\" (45 words)\n\n"
            "After (v8): \"Bas kat raha hu. Waha kaise chal raha hai?\" (9 words, Hinglish)\n\n"
            "Key params: temperature=0.8, top_p=0.85, repeat_penalty=1.2, max_tokens=120, "
            "random seed per call, 4-sentence hard truncation in post-processor."
        ),
    },
    {
        "day": "Day 3",
        "title": "Privacy Audit & Hardcoded Path Removal",
        "body": (
            "Found 3 files with hardcoded local filesystem paths "
            "(/Volumes/wininstall/...) that would leak machine layout if repo went public.\n\n"
            "Fixed:\n"
            "  \u2022 app.js: model paths \u2192 empty defaults (server discovers)\n"
            "  \u2022 main.rs: fallback path \u2192 current working directory\n"
            "  \u2022 smoke.py: hardcoded paths \u2192 env vars + auto-search\n\n"
            "Model discovery chain: env var \u2192 ~/.pratibmb/models/ \u2192 ~/models/ \u2192 ./models/\n\n"
            "Verified: no tokens, passwords, API keys, or personal data in git history. "
            "Commit emails use placeholder addresses."
        ),
    },
    {
        "day": "Day 3",
        "title": "Phase 1: Structured Profile Extraction",
        "body": (
            "The biggest intelligence upgrade. Used Gemma-3-4B itself to batch-analyze "
            "the entire corpus and build a structured identity profile.\n\n"
            "Extraction pipeline (~7 minutes on M-series Mac):\n"
            "  \u2022 89 threads analyzed for relationship type, topics, and summary\n"
            "  \u2022 Life events extracted per year (career, moves, travel, relationships)\n"
            "  \u2022 13 year summaries generated\n"
            "  \u2022 Communication style quantified\n\n"
            "Examples from real extraction:\n"
            "  \u2022 \"Rosaline Patra (former_romantic, talk about social_media, life_updates)\"\n"
            "  \u2022 \"[2014] Deeply involved in an internship at Samsung\"\n"
            "  \u2022 \"[2015] Traveling and working abroad, specifically in Rio\"\n"
            "  \u2022 \"[2023] Lost baggage during flight to Singapore\"\n\n"
            "Profile stored as JSON in SQLite. Context builder assembles a "
            "~200-token year-specific identity block for each chat turn."
        ),
    },
    {
        "day": "Day 3",
        "title": "Thread-Aware Retrieval",
        "body": (
            "Before: retrieved 8 isolated sentences by cosine similarity.\n"
            "After: each retrieved message includes 3 messages before and 3 after "
            "from the same thread.\n\n"
            "Messages grouped by thread in the prompt for coherent reading. "
            "The model now sees actual conversation flow:\n\n"
            "  [Thread: Sarah]\n"
            "    Sarah: hey how's the new place?\n"
            "    You: its good yaar, still unpacking\n"
            "    Sarah: lol typical\n\n"
            "Instead of:\n"
            "  [2016-03-15] its good yaar, still unpacking\n\n"
            "This lets the model understand context, relationships, and tone "
            "from how the user actually conversed."
        ),
    },
    {
        "day": "Day 3",
        "title": "Profile-Aware Chat \u2014 The Breakthrough",
        "body": (
            "System prompt now includes a year-specific identity block:\n\n"
            "  \"You are Tapas Kar, year 2015.\n"
            "  Style: you mix Hindi and English (Hinglish); your messages are short.\n"
            "  Life in 2015: Traveling and working abroad, specifically in Rio.\n"
            "  People: Rosaline Patra (former_romantic), Saini Aman (colleague)...\"\n\n"
            "Results with real corpus:\n\n"
            "  Q (2015): \"kya kar raha hai aajkal?\"\n"
            "  A: \"Bas kat raha hu. Waha kaise chal raha hai?\"\n\n"
            "  Q (2019): \"do you remember any fun times recently?\"\n"
            "  A: \"Yeah, I was just remembering Holi last month. We went around "
            "looking for good colors \u2013 it was quite an adventure.\"\n\n"
            "  Q (2013): \"who do you hang out with the most?\"\n"
            "  A: \"With my sofa. \U0001f60a\"\n\n"
            "Zero hallucination. Real names, real events, correct language per era."
        ),
    },
    {
        "day": "Day 3\u20134",
        "title": "Phase 2: LoRA Fine-Tuning Pipeline",
        "body": (
            "Built a complete fine-tuning pipeline to internalize the user's voice:\n\n"
            "1. Pair extraction: Find consecutive (other\u2192self) message pairs in threads. "
            "Include 2\u20133 messages of context. ~2,500 pairs from the corpus.\n\n"
            "2. Gemma formatting: Convert to Gemma chat template with "
            "<start_of_turn>user / <start_of_turn>model tags. Save as JSONL.\n\n"
            "3. Training: MLX-LM wrapper for native Metal acceleration. "
            "LoRA rank 16, alpha 32, targeting q/k/v/o projection layers. "
            "2\u20133 epochs, ~1 hour on M-series Mac.\n\n"
            "4. Conversion: MLX adapter \u2192 GGUF LoRA format for llama-cpp-python.\n\n"
            "The adapter (~10 MB) loads at inference via Llama(lora_path=...) with "
            "no changes to the base model file. Auto-detected from ~/.pratibmb/models/adapter.gguf."
        ),
    },
    {
        "day": "",
        "title": "Architecture Overview",
        "body": (
            "The full pipeline:\n\n"
            "  Import: User messages \u2192 Importers (WA/FB/IG) \u2192 SQLite corpus\n"
            "  Embed: Corpus \u2192 nomic-embed-text \u2192 float32 vectors in SQLite\n"
            "  Profile: Corpus \u2192 LLM batch analysis \u2192 relationships, events, summaries\n"
            "  LoRA: Corpus \u2192 training pairs \u2192 MLX-LM \u2192 adapter.gguf\n\n"
            "  Chat turn:\n"
            "    Query \u2192 Embed \u2192 Cosine retrieval (thread-aware)\n"
            "    + Profile context (year-specific identity)\n"
            "    + LoRA adapter (voice internalization)\n"
            "    \u2192 Gemma-3-4B generates response\n"
            "    \u2192 Post-processor cleans output\n\n"
            "All local. No cloud. No telemetry. 100% privacy.\n\n"
            "Stack: Python core + Tauri 2.x desktop + llama-cpp-python + Metal\n"
            "Total model size: 2.3 GB (chat) + 84 MB (embed) + ~10 MB (LoRA adapter)"
        ),
    },
    {
        "day": "",
        "title": "Metrics & Summary (Day 1\u20134)",
        "body": (
            "Corpus: 11,676 messages, 4,009 self-authored, 2010\u20132026\n"
            "Profile: 89 relationships, 8 life events, 13 year summaries\n"
            "Training pairs: ~2,500 conversation pairs for LoRA\n"
            "Tests: 22 unit tests passing\n"
            "Extraction time: ~7 min (profile), ~90 sec (embeddings)\n"
            "Chat latency: 2\u20134 sec per reply on Metal\n"
            "App size: ~3.9 MB (Tauri desktop, excl. models)\n\n"
            "Models:\n"
            "  \u2022 Gemma-3-4B-Instruct Q4_K_M (2.3 GB) \u2014 chat\n"
            "  \u2022 nomic-embed-text-v1.5 Q4_K_M (84 MB) \u2014 embeddings\n\n"
            "Repo: github.com/tapaskar/Pratibmb (private)\n"
            "License: AGPL v3"
        ),
    },
    {
        "day": "Day 5",
        "title": "Distribution Strategy \u2014 Package Manager Repos",
        "body": (
            "Evaluated two approaches for distributing the desktop app:\n\n"
            "  Option A: Separate repos per platform (Pratibmb-mac, -windows, -linux)\n"
            "  Option B: Monorepo + platform-specific distribution repos\n\n"
            "Decision: Option B. The codebase is 99% shared (Python backend, HTML/JS "
            "frontend, Rust shell). Splitting source code per platform would create "
            "sync nightmares. Every major Tauri/Electron app (VS Code, Signal, Obsidian) "
            "uses a monorepo.\n\n"
            "Instead, created thin distribution repos for package managers:\n"
            "  \u2022 homebrew-pratibmb \u2014 Homebrew Cask formula (brew install --cask pratibmb)\n"
            "  \u2022 pratibmb-aur \u2014 AUR package (yay -S pratibmb-bin)\n"
            "  \u2022 pratibmb-winget \u2014 winget manifest (winget install tapaskar.Pratibmb)\n\n"
            "Each is ~1 file, auto-updated by CI when a release is published. "
            "A new workflow (publish-packages.yml) downloads release assets, computes "
            "SHA256 hashes, and pushes updated manifests to each distribution repo."
        ),
    },
    {
        "day": "Day 5",
        "title": "Landing Page Overhaul",
        "body": (
            "Redesigned the downloads section on pratibmb.com:\n\n"
            "  Before: 4 cards (macOS ARM, macOS Intel, Linux, Windows) with direct download links\n"
            "  After: 3 cards (macOS, Linux, Windows) each showing:\n"
            "    1. Package manager one-liner (brew/yay/winget) as primary install\n"
            "    2. Direct download links as fallback\n"
            "    3. Install size with fine-tuning add-on info\n\n"
            "macOS card: '~450 MB install \u00b7 includes MLX fine-tuning'\n"
            "Linux/Windows: '~150 MB install \u00b7 fine-tuning is a one-command add-on (+3\u20135 GB)'\n\n"
            "Updated fine-tuning table: macOS Apple Silicon now shows 'Included' (green) "
            "instead of a pip command. MLX-LM added as default dependency in pyproject.toml "
            "with platform gate: sys_platform == 'darwin' and platform_machine == 'arm64'."
        ),
    },
    {
        "day": "Day 5",
        "title": "License \u2014 Dual AGPL + Commercial",
        "body": (
            "Switched from pure AGPLv3 to dual licensing:\n\n"
            "  \u2022 Open source: AGPLv3 (unchanged LICENSE file)\n"
            "  \u2022 Commercial: separate license for companies that can't comply with AGPL\n\n"
            "Rationale: AGPL alone lets cloud providers build hosted clones. "
            "Dual licensing keeps Pratibmb free for individuals and researchers "
            "while ensuring commercial use contributes back \u2014 either in code or funding.\n\n"
            "Model: Qt, MySQL, GitLab all use this pattern successfully. "
            "A CLA (Contributor License Agreement) is required for contributions "
            "so the project can offer commercial licenses without per-contributor permission.\n\n"
            "Updated: README, landing page hero badge, privacy section, footer, "
            "and all 3 distribution repo READMEs."
        ),
    },
    {
        "day": "Day 5",
        "title": "Desktop Help Modal \u2014 Navigation Paths + Troubleshooting",
        "body": (
            "Enhanced the in-app help (? button) with two additions:\n\n"
            "1. Visual navigation breadcrumbs for all 8 platforms:\n"
            "   Each section now shows a styled path indicator above the step list.\n"
            "   Example: [Open chat] \u2192 [Menu \u22ee] \u2192 [More] \u2192 [Export Chat] \u2192 [Without Media]\n"
            "   First step is accent-colored, rest are neutral pills with arrow separators.\n\n"
            "2. Troubleshooting FAQ (11 collapsible items):\n"
            "   3 general (no messages, wrong attribution, garbled text)\n"
            "   8 platform-specific (one per platform, covering the #1 gotcha each)\n"
            "   Uses native <details>/<summary> for zero-JS collapsible behavior.\n\n"
            "Design choice: breadcrumb paths instead of screenshots. Screenshots of "
            "third-party apps (WhatsApp, Facebook, etc.) are hard to maintain, "
            "add binary bulk, and may have copyright issues. Breadcrumbs are "
            "always accurate and match the app's design language."
        ),
    },
    {
        "day": "Day 5",
        "title": "Build Fixes \u2014 Windows and macOS Intel",
        "body": (
            "Diagnosed and fixed two CI build failures:\n\n"
            "1. Windows: icon.ico was 16x16 only (761 bytes). Tauri's Windows Resource "
            "generator requires multi-resolution ICO for the .exe manifest. "
            "Regenerated from icon.png (512x512) using ImageMagick to produce "
            "a proper 6-resolution ICO (256/128/64/48/32/16). New file: 130 KB.\n\n"
            "2. macOS Intel: macos-13 runner was deprecated by GitHub Actions. "
            "Jobs were cancelled instantly (0 seconds, 0 steps). Fix: cross-compile "
            "x86_64-apple-darwin target from macos-latest (ARM runner).\n\n"
            "Result: all 4 platform builds now pass. 6 assets uploaded to v0.1.0:\n"
            "  \u2022 Pratibmb_0.0.1_aarch64.dmg (4.7 MB)\n"
            "  \u2022 Pratibmb_0.0.1_x64.dmg (5.2 MB)\n"
            "  \u2022 Pratibmb_0.0.1_amd64.deb (5 MB)\n"
            "  \u2022 Pratibmb_0.0.1_amd64.AppImage (80.5 MB)\n"
            "  \u2022 Pratibmb_0.0.1_x64-setup.exe (3.1 MB)\n"
            "  \u2022 Pratibmb_0.0.1_x64_en-US.msi (4.8 MB)"
        ),
    },
    {
        "day": "Day 5",
        "title": "Security Audit \u2014 Going Public",
        "body": (
            "Full security audit before making the repo public:\n\n"
            "  \u2022 No API keys, tokens, or secrets in any tracked files\n"
            "  \u2022 No .env, .db, .sqlite, .jsonl, or .pem files in git history\n"
            "  \u2022 CI uses ${{ secrets.* }} properly (no hardcoded tokens)\n"
            "  \u2022 TAURI_SIGNING_PRIVATE_KEY set to empty string in CI\n"
            "  \u2022 Git history clean \u2014 no sensitive files ever committed then deleted\n"
            "  \u2022 Repo size: 1.9 MB (no large binaries in git)\n"
            "  \u2022 Commit emails use placeholder addresses\n\n"
            "One fix: added fused_model/ to .gitignore (was untracked but not ignored; "
            "contains merged LoRA weights with potential personal data baked in).\n\n"
            "Repo made PUBLIC. All 6 download links verified HTTP 200. "
            "Landing page live at pratibmb.com."
        ),
    },
    {
        "day": "Day 5",
        "title": "Model Distribution \u2014 The Missing Piece",
        "body": (
            "Discovered a critical gap: the README says 'models downloaded on first "
            "launch' but NO download code exists. Users would hit 'model not found' "
            "on first launch.\n\n"
            "Investigation:\n"
            "  \u2022 App uses GGUF format via llama-cpp-python (not PyTorch)\n"
            "  \u2022 Models must be pre-placed at ~/.pratibmb/models/\n"
            "  \u2022 server.py searches env var \u2192 ~/.pratibmb/models/ \u2192 ~/models/ \u2192 ./models/\n"
            "  \u2022 Pre-quantized GGUF files DO exist on HuggingFace (no quantization needed)\n\n"
            "Available on HuggingFace (public, no auth required):\n"
            "  \u2022 bartowski/google_gemma-3-4b-it-GGUF \u2192 Q4_K_M (2.49 GB)\n"
            "  \u2022 nomic-ai/nomic-embed-text-v1.5-GGUF \u2192 Q4_K_M (84 MB)\n\n"
            "Plan: Add huggingface_hub as a dependency. Implement auto-download in "
            "server.py's _env_model() \u2014 if model not found locally, download from "
            "HuggingFace with progress reporting. Show progress in onboarding wizard. "
            "Works offline after first download.\n\n"
            "This is the standard pattern (Ollama, LM Studio, GPT4All all do this)."
        ),
    },
    {
        "day": "",
        "title": "Deployment Architecture",
        "body": (
            "Full deployment architecture documented in docs/ARCHITECTURE.md:\n\n"
            "  User's Machine:\n"
            "    Tauri Desktop (.dmg/.exe/.deb/.AppImage)\n"
            "      \u2502\n"
            "      \u251c\u2500 WebView (HTML/JS/CSS) \u2190\u2192 Rust Backend (IPC)\n"
            "      \u2502                                \u2502\n"
            "      \u2502                    HTTP 127.0.0.1:11435\n"
            "      \u2502                                \u2502\n"
            "      \u251c\u2500 Python Server (pratibmb.server)\n"
            "      \u2502   \u251c\u2500 Importers (8 platforms)\n"
            "      \u2502   \u251c\u2500 RAG Engine (embed + cosine)\n"
            "      \u2502   \u251c\u2500 LLM Chat (Gemma-3-4B via llama.cpp)\n"
            "      \u2502   \u2514\u2500 Fine-tuning (MLX or PyTorch)\n"
            "      \u2502\n"
            "      \u251c\u2500 llama-cpp-python (Metal / CUDA / CPU)\n"
            "      \u2502   \u251c\u2500 Chat: gemma-3-4b-it Q4_K_M (2.3 GB)\n"
            "      \u2502   \u251c\u2500 Embed: nomic-embed-text-v1.5 Q4_K_M (84 MB)\n"
            "      \u2502   \u2514\u2500 LoRA adapter (~10 MB, optional)\n"
            "      \u2502\n"
            "      \u2514\u2500 SQLite Store (messages, embeddings, profile)\n\n"
            "  One-time download from HuggingFace Hub (first launch):\n"
            "    \u2022 bartowski/google_gemma-3-4b-it-GGUF (2.49 GB)\n"
            "    \u2022 nomic-ai/nomic-embed-text-v1.5-GGUF (84 MB)\n\n"
            "  Distribution:\n"
            "    \u2022 pratibmb.com (landing page, GitHub Pages)\n"
            "    \u2022 Homebrew Cask (macOS)\n"
            "    \u2022 AUR (Arch Linux)\n"
            "    \u2022 winget (Windows)\n"
            "    \u2022 Direct downloads (.dmg, .deb, .AppImage, .exe, .msi)"
        ),
    },
    {
        "day": "",
        "title": "Metrics & Summary (Updated)",
        "body": (
            "Corpus: 11,676 messages, 4,009 self-authored, 2010\u20132026\n"
            "Profile: 89 relationships, 8 life events, 13 year summaries\n"
            "Training pairs: ~2,500 conversation pairs for LoRA\n"
            "Tests: 22 unit tests passing\n"
            "Extraction time: ~7 min (profile), ~90 sec (embeddings)\n"
            "Chat latency: 2\u20134 sec per reply on Metal\n\n"
            "App size: 3\u20135 MB (Tauri desktop, excl. models)\n"
            "Model download: ~2.6 GB one-time (first launch)\n"
            "Install size: ~450 MB (macOS w/ MLX), ~150 MB (Windows/Linux base)\n\n"
            "Models:\n"
            "  \u2022 Gemma-3-4B-Instruct Q4_K_M (2.3 GB) \u2014 chat\n"
            "  \u2022 nomic-embed-text-v1.5 Q4_K_M (84 MB) \u2014 embeddings\n\n"
            "Platforms: macOS (ARM+Intel), Linux (deb+AppImage), Windows (exe+msi)\n"
            "Install methods: brew, yay, winget, direct download\n"
            "Repo: github.com/tapaskar/Pratibmb (public)\n"
            "License: AGPL v3 + Commercial dual license"
        ),
    },
]


def build():
    doc = Document()

    # ---- Styles ----
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Arial"
        h.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

    # ---- Title page ----
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Pratibmb")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xB0, 0x89, 0x68)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Journal\nBuilding a Local AI That Talks Like You")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("\n\n100% local \u00b7 no cloud \u00b7 no telemetry")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ---- Table of Contents placeholder ----
    doc.add_heading("Table of Contents", level=1)
    for i, entry in enumerate(ENTRIES, 1):
        toc_line = doc.add_paragraph()
        day_prefix = f"{entry['day']} \u2014 " if entry["day"] else ""
        run = toc_line.add_run(f"{i}. {day_prefix}{entry['title']}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ---- Entries ----
    for i, entry in enumerate(ENTRIES, 1):
        day_prefix = f"{entry['day']} \u2014 " if entry["day"] else ""
        doc.add_heading(f"{i}. {day_prefix}{entry['title']}", level=2)

        for para_text in entry["body"].split("\n\n"):
            p = doc.add_paragraph()
            # Handle indented lines as a block
            lines = para_text.split("\n")
            for j, line in enumerate(lines):
                if j > 0:
                    p.add_run("\n")
                run = p.add_run(line)
                # Style indented lines slightly differently
                if line.startswith("  "):
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build()
